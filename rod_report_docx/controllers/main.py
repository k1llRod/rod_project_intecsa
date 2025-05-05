import os
from odoo import http
from odoo.http import request
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt, Inches
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from num2words import num2words
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime


import os
import base64
import io
class SaleOrderReportController(http.Controller):

    @http.route('/sale_order/report/<int:order_id>', type='http', auth='user', website=True)
    def get_report(self, order_id):
        sale_order = request.env['sale.order'].browse(order_id)
        env = request.env
        if not sale_order.exists():
            return request.not_found()

        # Generar el documento Word
        document = Document()
        # document.add_picture(image_stream, width=Inches(2))
        # Agregar encabezado con imagen
        section = document.sections[0]

        section.left_margin = Inches(0.4) #margen de hoja
        section.right_margin = Inches(0.4)
        header = section.header
        header_paragraph = header.paragraphs[0]
        module_path = os.path.dirname(os.path.abspath(__file__))
        # Ruta de la imagen (reemplazar con la ruta correcta o usar un stream si la imagen viene de la base de datos)
        # image_path = "custom_addons/rod_project_intecsa/rod_report_docx/static/src/img/intecsa.jpeg"
        # image_path_footer = "custom_addons/rod_project_intecsa/rod_report_docx/static/src/img/footer.png"
        image_path = os.path.join(module_path, '..', 'static', 'src', 'img', 'intecsa.jpeg')
        image_path_footer = os.path.join(module_path, '..', 'static', 'src', 'img', 'footer.png')
        header_paragraph.add_run().add_picture(image_path, width=Inches(1.25))
        # header_paragraph.add_run(f'Orden de entrega: {sale_order.name}')

        # Agregar pie de página
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.add_run().add_picture(image_path_footer, width=Inches(7.5))


        # document.add_heading('Orden de entrega', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        for record in sale_order:
            paragraph_code = document.add_paragraph()
            paragraph_code.paragraph_format.space_before = Pt(0)
            paragraph_code.paragraph_format.space_after = Pt(0)
            run = paragraph_code.add_run(record.name)
            paragraph_code.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in record.picking_ids:
                paragraph_code = document.add_paragraph()
                paragraph_code.paragraph_format.space_before = Pt(0)
                paragraph_code.paragraph_format.space_after = Pt(0)
                run = paragraph_code.add_run(r.name_ent)
                paragraph_code.alignment = WD_ALIGN_PARAGRAPH.RIGHT


        table_information = document.add_table(rows=1, cols=2)
        table_information.style = 'Table Grid'
        table_information.autofit = False
        first_row = table_information.rows[0].cells
        first_row[0].width = Inches(2)  # 25% de 6 pulgadas
        first_row[1].width = Inches(4)  # 75% de 6 pulgadas

        # Fusionamos la primera celda con la segunda
        merged_cell = first_row[0].merge(first_row[1])

        # Agregamos el texto en la celda fusionada
        paragraph = merged_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run('NOTA DE ENTREGA').bold = True


        for rec in sale_order:


            hdr_cells = table_information.add_row().cells
            # Nombre
            hdr_cells[0].width = 10  # 25% de 6 pulgadas
            hdr_cells[1].width = 20  # 75% de 6 pulgadas
            hdr_cells[0].paragraphs[0].add_run('Nombre').bold = True
            hdr_cells[1].text = rec.partner_id.name
            hdr_cells[0].width = 100

            # NIT
            hdr_cells = table_information.add_row().cells
            hdr_cells[0].paragraphs[0].add_run('NIT').bold = True
            hdr_cells[1].text = rec.partner_id.vat if rec.partner_id.vat else ''

            # Dirección
            hdr_cells = table_information.add_row().cells
            hdr_cells[0].paragraphs[0].add_run('Dirección').bold = True
            hdr_cells[1].text = rec.partner_id.street if rec.partner_id.street else ''

            # Telefono
            hdr_cells = table_information.add_row().cells
            hdr_cells[0].paragraphs[0].add_run('Telefono').bold = True
            hdr_cells[1].text = rec.partner_id.mobile if rec.partner_id.mobile else ''

            # Departamento
            hdr_cells = table_information.add_row().cells
            hdr_cells[0].paragraphs[0].add_run('Departamento').bold = True
            hdr_cells[1].text = rec.partner_id.state_id.name if rec.partner_id.state_id.name else ''

            # Forma de pago
            hdr_cells = table_information.add_row().cells
            hdr_cells[0].paragraphs[0].add_run('Forma de pago').bold = True
            method =  rec.payment_method if rec.payment_method else ''
            if method == 'transfer':
                hdr_cells[1].text = 'Transferencia'
            elif method == 'transfer_sigep':
                hdr_cells[1].text = 'Transferencia SIGEP'
            elif method == 'cheque':
                hdr_cells[1].text = 'Cheque'
            elif method == 'qr':
                hdr_cells[1].text = 'QR'

            # Fecha de entrega
            hdr_cells = table_information.add_row().cells
            hdr_cells[0].paragraphs[0].add_run('Fecha de entrega').bold = True
            hdr_cells[1].text = rec.delivery_text if rec.delivery_text else ''
            # hdr_cells = table_information.add_row().cells

        from docx.enum.table import WD_TABLE_ALIGNMENT
        for row in table_information.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
                # for paragraph in cell.paragraphs:
                #     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph('')

        tbl1 = table_information._tbl
        for cell in tbl1.iter_tcs():
            tc_pr = cell.tcPr
            borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'</w:tcBorders>')
            tc_pr.append(borders)


        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        # hdr_cells = table.add_row().cells
        # hdr_cells = table.add_column(width=10)

        # Asignar texto y ancho a cada columna (por ejemplo: total 6.5 pulgadas)
        hdr_cells[0].width = Inches(2)  # Descripción
        hdr_cells[1].width = Inches(1)  # Cantidad
        hdr_cells[2].width = Inches(1)  # Precio unitario
        hdr_cells[3].width = Inches(1)  # Precio total

        headers = ['Descripción', 'Cantidad', 'Precio unitario', 'Precio total']
        for i, cell in enumerate(hdr_cells):
            # Texto
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(headers[i])
            run.bold = True

            # Alineación vertical al centro
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Fondo gris plomo (transparente claro)
            cell._tc.get_or_add_tcPr().append(parse_xml(
                r'<w:shd {} w:fill="bfbfbf "/>'.format(nsdecls('w'))
            ))

        hdr_cells[0].text = 'Descripcion'
        hdr_cells[1].text = 'Cantidad'
        hdr_cells[2].text = 'Precio unitario'
        hdr_cells[3].text = 'Precio total'

        for line in sale_order.order_line:
            row_cells = table.add_row().cells
            if line.display_type == False:
                row_cells[0].text = line.product_id.name
                row_cells[1].text = str(line.product_uom_qty)
                row_cells[2].text = str(line.price_unit)
                row_cells[3].text = str(line.price_total)
            if line.display_type == 'line_note':
                # row_cells[0].merge(row_cells[3])
                row_cells[0].text = line.name
                # row_cells[1].merge(row_cells[2])
                # row_cells[1].text = ''

        row_cells = table.add_row().cells
        row_cells[0].merge(row_cells[2])
        row_cells[0].paragraphs[0].add_run('Total').bold = True
        row_cells[3].text = str(sale_order.amount_total) if sale_order.amount_total else ''

        literal_number = num2words(int(sale_order.amount_total), lang='es').upper()
        decimal = str(round(sale_order.amount_total % 1 * 100))
        literal_number = literal_number + ', CON ' + decimal + '/100 BOLIVIANOS'

        row_cells = table.add_row().cells
        row_cells[0].merge(row_cells[3])
        row_cells[0].paragraphs[0].add_run('Precio total: ' + literal_number).bold = True
        # row_cells[3].text = str(sale_order.amount_total) if sale_order.amount_total else ''



        tbl = table._tbl
        for cell in tbl.iter_tcs():
            tc_pr = cell.tcPr
            borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'</w:tcBorders>')
            tc_pr.append(borders)

        document.add_paragraph('')
        document.add_heading('OBSERVACIONES', level=1)
        table_firmas = document.add_table(rows=1, cols=3)
        table_firmas.style = 'Table Grid'
        # tbl2 = table_firmas._tbl
        for record in sale_order:
            hdr_cells = table_firmas.rows[0].cells
            hdr_cells[0].text = '\n' + '\n' + '\n' + '\n'
            hdr_cells[1].text = '\n' + '\n' + '\n' + '\n'
            hdr_cells[2].text = '\n' + '\n' + '\n' + '\n'
            row_cells = table_firmas.add_row().cells
            row_cells[0].text = 'MANAGER OFFICE'
            paragraph = row_cells[0].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].text = 'ASESOR COMERCIAL'
            paragraph = row_cells[1].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].text = 'CLIENTE'
            paragraph = row_cells[2].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells = table_firmas.add_row().cells
            row_cells[0].text = '\n' + 'Lic. Mario Robles M'
            paragraph = row_cells[0].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].text =  '\n' + record.user_id.name
            paragraph = row_cells[1].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].text = '\n' + 'Cargo:_________________________' + '\n' + 'Nombre:_________________________' + '\n'
            paragraph = row_cells[2].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # row_cells = table_firmas.add_row().cells

        tbl2 = table_firmas._tbl
        for cell in tbl2.iter_tcs():
            tc_pr = cell.tcPr
            borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'</w:tcBorders>')
            tc_pr.append(borders)

        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        file_data = file_stream.read()

        # Crear y devolver la respuesta HTTP
        return request.make_response(
            file_data,
            headers=[
                ('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
                ('Content-Disposition', f'attachment; filename=Sales_Order_{sale_order.name}.docx')
            ]
        )
    @http.route('/sale_order/report_warranty/<int:order_id>', type='http', auth='user', website=True)
    def get_report_warranty(self, order_id):
        sale_order = request.env['sale.order'].browse(order_id)
        env = request.env
        if not sale_order.exists():
            return request.not_found()

        # Generar el documento Word
        document = Document()
        # document.add_picture(image_stream, width=Inches(2))
        # Agregar encabezado con imagen
        section = document.sections[0]

        section.left_margin = Inches(0.4)  # margen de hoja
        section.right_margin = Inches(0.4)
        header = section.header
        header_paragraph = header.paragraphs[0]
        module_path = os.path.dirname(os.path.abspath(__file__))
        # Ruta de la imagen (reemplazar con la ruta correcta o usar un stream si la imagen viene de la base de datos)
        # image_path = "custom_addons/rod_project_intecsa/rod_report_docx/static/src/img/intecsa.jpeg"
        # image_path_footer = "custom_addons/rod_project_intecsa/rod_report_docx/static/src/img/footer.png"
        image_path = os.path.join(module_path, '..', 'static', 'src', 'img', 'intecsa.jpeg')
        image_path_footer = os.path.join(module_path, '..', 'static', 'src', 'img', 'footer.png')
        header_paragraph.add_run().add_picture(image_path, width=Inches(1.25))
        # header_paragraph.add_run(f'Orden de entrega: {sale_order.name}')

        # Agregar pie de página
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.add_run().add_picture(image_path_footer, width=Inches(7.5))

        document.add_heading('CERTIFICADO DE GARANTIA', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Crear una tabla de 1 fila y 3 columnas
        table_header = document.add_table(rows=1, cols=3)
        table_header.autofit = False
        table_header.allow_autofit = False
        table_header.style = 'Table Grid'

        # Ajustar anchos de columna
        table_header.columns[0].width = Inches(2.5)
        table_header.columns[1].width = Inches(3)
        table_header.columns[2].width = Inches(2)

        for record in sale_order:
            cell_left = table_header.cell(0, 0).paragraphs[0]
            run_left = cell_left.add_run("N°. " + record.name+'\n')
            run_left.bold = True
            cell_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in record.picking_ids:
                cell_left = table_header.cell(0, 0).paragraphs[0]
                run_left = cell_left.add_run(r.name_ent + '\n')
                run_left.bold = True
                cell_left.alignment = WD_ALIGN_PARAGRAPH.LEFT

        cell_right = table_header.cell(0, 2).paragraphs[0]
        run_right = cell_right.add_run(datetime.now().strftime("%d/%m/%Y"))
        run_right.bold = True
        cell_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        tbl3 = table_header._tbl
        for cell in tbl3.iter_tcs():
            tc_pr_header = cell.tcPr
            borders = parse_xml(
                r'<w:tcBorders %s>'
                r'<w:top w:val="nil"/>'
                r'<w:left w:val="nil"/>'
                r'<w:bottom w:val="nil"/>'
                r'<w:right w:val="nil"/>'
                r'</w:tcBorders>' % nsdecls('w')
            )
            tc_pr_header.append(borders)




        # Agregar un párrafo personalizado
        intro_paragraph = document.add_paragraph()
        intro_paragraph.paragraph_format.space_before = Pt(12)
        intro_paragraph.paragraph_format.space_after = Pt(12)
        intro_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        intro_paragraph.add_run(
            "La empresa INTECSA SRL, en todos sus productos nuevos tiene la garantía correspondiente "
            "del fabricante por el plazo que indica cada producto, salvo algunas excepciones claramente"
            "identificadas. El fabricante de la marca es el único responsable de esta garantía y la"
            "función de INTECSA SRL. es la de tramitar esa garantía ante el fabricante, por medio"
            "de su Centro Autorizado de Servicio (CAS), para su respectiva reparación y/o cambio de"
            "la parte dañada. INTECSA SRL., hará todo lo posible para procesar todos los productos"
            "recibidos para reparación por garantía en el menor tiempo posible. En este contexto,"
            "se establece que INTECSA SRL estará sujeto a los tiempos de ejecución de garantía propia"
            "de la marca, para mayor información, escanear el código QR y/o visita: https://intecsa.com.bo/politicas/"
            "Los equipos en garantia son:"
        )


        section = document.sections[0]
        section.top_margin = Pt(20)
        section.bottom_margin = Pt(20)
        section.left_margin = Pt(20)
        section.right_margin = Pt(20)

        # Crear tabla principal
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False

        # Tamaños de las columnas
        table.columns[0].width = Inches(5.2)
        table.columns[1].width = Inches(2.0)

        # Celda izquierda: productos
        cell_left = table.cell(0, 0)
        p = cell_left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Listar productos
        for line in sale_order.order_line:
            if not line.display_type:
                p.add_run(f"{int(line.product_uom_qty)} {line.product_id.name}\n").bold = False

        # Texto adicional
        p.add_run("\n" + sale_order.warranty +"DE FÁBRICA ORIGINAL\n").bold = True
        p.add_run("COMSURRAL XXI\n").bold = True
        p.add_run(f"SEGÚN DETALLE DE CARACTERÍSTICAS EN NOTA DE ENTREGA N.º {sale_order.picking_ids[-1].name_ent}\n").bold = True

        # Celda derecha: imagen y título
        cell_right = table.cell(0, 1)
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_right.add_run("POLÍTICAS DE GARANTÍA\n").bold = True

        # Agregar imagen (QR u otro)
        image_path = os.path.join(os.path.dirname(__file__),'..', 'static', 'src', 'img', 'qr_warranty.png')  # actualiza ruta
        if os.path.exists(image_path):
            p_right.add_run().add_picture(image_path, width=Inches(1.5))

        tbl = table._tbl
        for cell in tbl.iter_tcs():
            tc_pr = cell.tcPr
            borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'</w:tcBorders>')
            tc_pr.append(borders)

        # Texto introductorio
        intro = "Cualquier defecto debido a componentes o materiales, está cubierto por la presente garantía, con la excepción de los siguientes casos:"
        p_intro = document.add_paragraph(intro)
        p_intro.paragraph_format.space_after = Pt(6)

        # Lista de excepciones con viñetas
        lista = [
            "Problemas derivados de transporte después de la compra.",
            "Falla producida por conexiones externas, derivadas de la instalación, conexión a fuentes de alimentación diferentes a las indicaciones de las especificaciones técnicas del producto.",
            "Mal manejo, mal trato, uso indebido del producto o uso de accesorios inadecuados.",
            "Uso del producto bajo condiciones externas: intemperie, humedad, alta temperatura.",
            "Intervención de personal ajeno a nuestros Servicios Técnicos Autorizados.",
            "Casos fortuitos o de causa mayor.",
            "Daños accidentales, o por negligencia del CLIENTE en el uso de enchufe, voltaje inadecuado u otra alteración que afecte la confiabilidad y calidad de la unidad no atribuible a una falla en la fabricación del equipo de computación y sus accesorios.",
            "Soporte técnico en laboratorio técnico del proveedor.",
        ]

        for item in lista:
            p = document.add_paragraph(item, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)

        # Texto de cierre
        cierre = (
            "De acuerdo a las políticas y condiciones de garantía de la empresa InteCSA SRL, "
            "el cliente firma en señal de conformidad:"
        )
        p_cierre = document.add_paragraph()
        p_cierre.add_run(cierre)
        p_cierre.paragraph_format.space_before = Pt(10)

        document.add_paragraph('')
        document.add_heading('OBSERVACIONES', level=1)
        table_firmas = document.add_table(rows=1, cols=3)
        table_firmas.style = 'Table Grid'
        # tbl2 = table_firmas._tbl
        for record in sale_order:
            hdr_cells = table_firmas.rows[0].cells
            hdr_cells[0].text = '\n' + '\n' + '\n' + '\n'
            hdr_cells[1].text = '\n' + '\n' + '\n' + '\n'
            hdr_cells[2].text = '\n' + '\n' + '\n' + '\n'
            row_cells = table_firmas.add_row().cells
            row_cells[0].text = 'MANAGER OFFICE'
            paragraph = row_cells[0].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].text = 'ASESOR COMERCIAL'
            paragraph = row_cells[1].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].text = 'CLIENTE'
            paragraph = row_cells[2].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells = table_firmas.add_row().cells
            row_cells[0].text = '\n' + 'Lic. Mario Robles M'
            paragraph = row_cells[0].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].text = '\n' + record.user_id.name
            paragraph = row_cells[1].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[
                2].text = '\n' + 'Cargo:_________________________' + '\n' + 'Nombre:_________________________' + '\n'
            paragraph = row_cells[2].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # row_cells = table_firmas.add_row().cells

        tbl2 = table_firmas._tbl
        for cell in tbl2.iter_tcs():
            tc_pr = cell.tcPr
            borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                r'</w:tcBorders>')
            tc_pr.append(borders)

        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        file_data = file_stream.read()

        # Crear y devolver la respuesta HTTP
        return request.make_response(
            file_data,
            headers=[
                ('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
                ('Content-Disposition', f'attachment; filename=Sales_Order_{sale_order.name}.docx')
            ]
        )

