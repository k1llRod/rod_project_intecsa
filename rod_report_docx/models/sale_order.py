from odoo import models, fields, api
from docx import Document
from io import BytesIO
import base64

class SaleOrderReport(models.Model):
    _inherit = 'sale.order'

    def generate_word_report(self):
        for order in self:
            document = Document()
            # Agregar encabezado
            header = document.sections[0].header
            header_paragraph = header.paragraphs[0]
            header_paragraph.text = f'Orden de entrega: {order.name}'

            # Agregar pie de página
            footer = document.sections[0].footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.text = 'Este es un pie de página.'

            document.add_heading('Orden de entrega', level=1)
            document.add_paragraph(f'Orden: {order.name}')
            document.add_paragraph(f'Fecha: {order.date_order}')
            document.add_paragraph(f'Cliente: {order.partner_id.name}')

            table = document.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Descripcion'
            hdr_cells[1].text = 'Cantidad'
            hdr_cells[2].text = 'Precio unitario'
            hdr_cells[2].text = 'Precio total'

            for line in order.order_line:
                row_cells = table.add_row().cells
                row_cells[0].text = line.product_id.name
                row_cells[1].text = str(line.product_uom_qty)
                row_cells[2].text = str(line.price_unit)
                row_cells[3].text = str(line.price_total)

            file_stream = BytesIO()
            document.save(file_stream)
            file_stream.seek(0)
            file_data = base64.b64encode(file_stream.read())

            attachment = self.env['ir.attachment'].create({
                'name': f'Sales_Order_{order.name}.docx',
                'type': 'binary',
                'datas': file_data,
                'res_model': 'sale.order',
                'res_id': order.id,
                'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            })
        return True